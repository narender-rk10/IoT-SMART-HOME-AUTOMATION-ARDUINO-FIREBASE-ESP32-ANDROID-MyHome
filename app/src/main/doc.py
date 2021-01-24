import docx
import os
import glob

#Creatinf a document
document = docx.Document()
font = document.styles["Normal"].font
font.name = "Times New Roman"
font.size = docx.shared.Pt(12)


#changes  the directory to the app
os.chdir("res/values")
#gets the current working directory
path = os.getcwd()
#iterates throught all the files recursively in all files and folders
for filename in glob.iglob( path + '**/**', recursive=True):
    #print(filename)
    if os.path.isdir(filename) == True:
        continue
    else:
        print(filename)
        #read the files recursivly from the folders
        code_file = open(filename,errors='ignore')
        file_data = code_file.read()
        #append the file data to the document
        file_title = os.path.basename(filename)
        document.add_heading(file_title)
        document.add_paragraph(file_data)
        document.add_page_break()
        #Close the file
        code_file.close()
        
document.save("CodeImplmentationValues.docx")
